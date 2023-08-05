# -*- coding: UTF-8 -*-
# 读取docx中的文本代码示例
import sys

import docx

from modules.scheduler.scheudler import RedisTaskScheduler
from repository import dao

reload(sys)  # Python2.5 初始化后删除了 sys.setdefaultencoding 方法，我们需要重新载入
sys.setdefaultencoding('utf-8')

task_scheduler = RedisTaskScheduler()

KEY_COVER_PRODUCTS_APPOINT = "farfetch:cover:product:appoint"

farfetch_urls = []
farfetch_codes = []


def create_urls_from_file(file_path):
    # 获取文档对象
    file = docx.Document(file_path)
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    for para in file.paragraphs:
        text = para.text
        if text.startswith("https"):
            farfetch_urls.append(text.strip())


# # 输出段落编号及段落内容
# for i in range(len(file.paragraphs)):
#     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)


if __name__ == '__main__':
    file_path = "data/xiaobao/julianfashion1.docx"
    create_urls_from_file(file_path)
    file_path = "data/xiaobao/wise1.docx"
    create_urls_from_file(file_path)

    for url in farfetch_urls:
        splices = url.split("-")
        last_splice = splices[-1]
        splices2 = last_splice.split(".")
        print splices2[0], "<--->", url
        farfetch_codes.append(splices2[0])

    print "total codes length:", len(farfetch_codes)
    print ",".join(farfetch_codes)

    for item in dao.common.get_appoint_products():
        task_scheduler.push(KEY_COVER_PRODUCTS_APPOINT, item["spider_product_id"])
