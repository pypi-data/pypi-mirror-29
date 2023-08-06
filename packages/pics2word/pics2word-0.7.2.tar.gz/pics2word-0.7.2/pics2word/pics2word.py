#!/usr/bin/env python3
import os, sys
from docx import Document
from docx.shared import Inches
import datetime, imghdr, struct
import math

class pics2word:

    def __init__(self):
        # set default values on instantiation, until changed with CL args
        self.SetPath()
        self.SetTitle() 
        self.SetPicWidth()
        self.SetPicHeight()
        self.GetPicsInPath()
        self.SetFormat()
        self.SetTableWidth()

    def SetPath(self, Path=os.getcwd()):
        # Default path is the current working directory on the command line
        self.path = Path

    def SetTitle(self, title="PhotoDoc", date='y'):
        # if date begions with 'y', append title with date
        if date[0] == 'y':
            Today = self.GetDate()
            self.title = title + "_" + str(Today)
        else:
            self.title = title

    def SetPicWidth(self,Width=4):
        # TODO set a default!
        self.width = Width

    def SetPicHeight(self,Height=4):
        # TODO set a default!
        self.height = Height

    def SetTableWidth(self, Columns=2):
        # TODO set a default!
        self.tablecolumns = Columns

    def GetDate(self):
        return datetime.date.today().strftime("%d%b%Y") # i.e. 15Feb2018

    def help(self):
        message = '''Usage: pics2word [-command] [value]
Options:
\t-h\t- Pass "help" to print this help message to the terminal. \n\t\t  Pass the name of a command below without the '-' for more informatio about that command. <UNDER CONSTRUCTION>
\t-P\t- Pass an alternative path to be used. i.e. \"C:\\\\Pictures\\\". Defaults to current directory.
\t-f\t- format pictures. pass either "normal" or "table". Defaults to normal. 
\t-T\t- Override the default title. Defaults to PhotoDoc_<current date> (See Td, below).
\t-Td\t- Choose to append the title with the current date. Options: \"y\" or \"n\". Defaults to \"y\".
\t-pw\t- Set the width of imported pictures in inches. Defaults to 4 inches
\t-ph\t- Set the height of imported pictures in inches. Defaults to 4 inches
\t-tw\t- Set the number of columns used in table format. Note: table format must be enabled! Defaults to 2.

Commands may be passed as command-value pairs in any order.
All commands are optional and the defaults will be used if no commands are given.

Example: pics2word -P \"C:\\\\Pictures\\\" -T \"Report\" -Td \"n\" -f \"table\"\n'''
        print(message)
    
    def SetFormat(self, format="normal"):
        if format[0].lower() == 't':
            self.format = "table"
        elif format[0].lower() == 'n':
            self.format = "normal"
        else:
            raise ValueError("Please enter a valid format for '-f' i.e. \"Normal\" or \"Table\"")

    def WriteDoc(self):
        if self.format[0].lower() == 't':
            self.WriteTable()
        else:
            self.WriteNormal()

    def WriteNormal(self):
        document = Document()
        p = document.add_paragraph()
        Path = self.path
        # Todo check if numbered and sort appropriately
        PicList = sorted(self.pics) # Sort pics into an order
        i=0
        for Pic in PicList:
            FullImageandPath = os.path.join(Path,Pic)
            r = p.add_run()
            isPortrait = self.IsPortrait(FullImageandPath)
            if isPortrait:
                r.add_picture(FullImageandPath,height=Inches(self.height))
            else:
                r.add_picture(FullImageandPath,width=Inches(self.width))
            p.add_run("\n"+Pic.split('.')[0]+"\n")
            # update progress
            cli_progress_test(cur_val=i,end_val=len(PicList))
            i += 1
        document.save(self.title + '.docx')

    def WriteTable(self):
        document = Document()
        Path = self.path
        PicList = sorted(self.pics) # Sort pics into an order
        numRows = self.GetNumberofRows()
        table = document.add_table(rows=numRows, cols=self.tablecolumns)
        i=0
        Row_Index = 0 # Resets every iteration
        # list[start:stop:step] pastes the picture in every 2nd cell
        for row in table.rows[::2]:
            Col_Index = 0 # Resets every iteration
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    try:
                        Pic = PicList[i]
                        FullImageandPath = os.path.join(Path,Pic)
                        r = paragraph.add_run()
                        isPortrait = self.IsPortrait(FullImageandPath)
                        if isPortrait:
                            r.add_picture(FullImageandPath,height=Inches(self.height))
                        else:
                            r.add_picture(FullImageandPath,width=Inches(self.width))
                        # Offset 1 row down and add description
                        table.cell(row_idx=Row_Index + 1,col_idx=Col_Index).text = Pic.split('.')[0]
                        # Update user of progress
                        cli_progress_test(cur_val=i,end_val=len(PicList))
                    except IndexError:
                        # we incur an index error at the end of the picture list
                        # hence, we will simply pass and do nothing with the remaining empty cells
                        pass
                Col_Index += 1
                i += 1
            Row_Index += 2
        document.save(self.title + '.docx')
    
    def GetPicsInPath(self):
        self.pics = []
        ValidExtList = [".jpg",".jpeg",".png",".bmp",".gif",".JPG",".JPEG",".PNG",".BMP",".GIF"]
        for file in os.listdir(self.path):
            for ValidExt in ValidExtList:
                if file.endswith(ValidExt):
                    self.pics.append(file)
        return self.pics

    def IsPortrait(self, fname):
        """Determine the image type of fhandle and return its size."""
        with open(fname, 'rb') as fhandle:
            head = fhandle.read(24)
            if len(head) != 24:
                return
            if imghdr.what(fname) == 'png':
                check = struct.unpack('>i', head[4:8])[0]
                if check != 0x0d0a1a0a:
                    return
                width, height = struct.unpack('>ii', head[16:24])
            elif imghdr.what(fname) == 'gif':
                width, height = struct.unpack('<HH', head[6:10])
            elif imghdr.what(fname) == 'jpeg':
                try:
                    fhandle.seek(0) # Read 0xff next
                    size = 2
                    ftype = 0
                    while not 0xc0 <= ftype <= 0xcf:
                        fhandle.seek(size, 1)
                        byte = fhandle.read(1)
                        while ord(byte) == 0xff:
                            byte = fhandle.read(1)
                        ftype = ord(byte)
                        size = struct.unpack('>H', fhandle.read(2))[0] - 2
                    # We are at a SOFn block
                    fhandle.seek(1, 1)  # Skip `precision' byte.
                    height, width = struct.unpack('>HH', fhandle.read(4))
                except Exception: #IGNORE:W0703
                    return
            else:
                return
            if width/height > 1:
                return False
            else:
                return True  

    def isNumbered(self,list):
        count = 0
        for value in list:
            string = value.split('.')[0]
            if string[-1].isdigit() or string[0].isdigit:
                count += 1
        # if all names start or end with numbers, 
        # then we can assume they have been numbered
        if count == len(list):
            sorting_tuple = [()]
            for value in list:
                string = value.split('.')[0]
                string[len(string.rstrip('0123456789')):]
            return True
        else:
            return False

    def GetNumberofRows(self):
        cols = self.tablecolumns
        NumofPics = len(self.pics)
        return int(math.ceil(NumofPics / cols)) * 2

def cli_progress_test(cur_val, end_val, bar_length=60, suffix=''):
    
    filled_len = int(round(bar_length * cur_val / float(end_val)))

    percents = round(100.0 * cur_val / float(end_val), 1)
    bar = '=' * filled_len + '-' * (bar_length - filled_len)

    sys.stdout.write('[%s] %s%s ...%s\r' % (bar, percents, '%', suffix))
    sys.stdout.flush()